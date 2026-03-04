import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document


def parse_args() -> argparse.Namespace:
    backend_dir = Path(__file__).resolve().parent
    project_root = backend_dir.parent
    workspace_root = project_root.parent
    reports_dir = workspace_root / "reports"

    parser = argparse.ArgumentParser(
        description="Generate DOCX summary report from consolidated JSON report."
    )
    parser.add_argument(
        "--input",
        default=str(reports_dir / "api_batch_all_models_20260303_224328.json"),
        help="Path to consolidated JSON report.",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output DOCX path. Default: reports/final_batch_report_<timestamp>.docx",
    )
    return parser.parse_args()


def load_report(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def format_accuracy(value: float | None) -> str:
    if value is None:
        return "N/A"
    return f"{value * 100:.1f}%"


def build_doc(report: dict, input_path: Path, output_path: Path) -> None:
    document = Document()
    document.add_heading("Raport zbiorczy - analiza phishing (batch)", level=1)

    generated_at = report.get("generated_at_utc", "-")
    api_url = report.get("api_url", "-")
    dataset_path = report.get("input_file", "-")
    models = report.get("models", [])
    model_reports = report.get("reports", [])

    document.add_paragraph(f"Data wygenerowania raportu źródłowego (UTC): {generated_at}")
    document.add_paragraph(f"API URL: {api_url}")
    document.add_paragraph(f"Dataset: {dataset_path}")
    document.add_paragraph(f"Liczba modeli: {len(models)}")

    document.add_heading("Podsumowanie wyników modeli", level=2)

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Model"
    headers[1].text = "Total"
    headers[2].text = "Success"
    headers[3].text = "Errors"
    headers[4].text = "Accuracy"
    headers[5].text = "Predykcje phishing"
    headers[6].text = "Predykcje legit"

    for report_row in model_reports:
        summary = report_row.get("summary", {})
        pred_dist = summary.get("prediction_distribution", {}) or {}

        row = table.add_row().cells
        row[0].text = str(report_row.get("model", "-"))
        row[1].text = str(summary.get("total", "-"))
        row[2].text = str(summary.get("success", "-"))
        row[3].text = str(summary.get("errors", "-"))
        row[4].text = format_accuracy(summary.get("accuracy"))
        row[5].text = str(pred_dist.get("phishing", 0))
        row[6].text = str(pred_dist.get("legit", 0))

    completed = [r for r in model_reports if (r.get("summary", {}).get("success", 0) or 0) > 0]
    failed = [r for r in model_reports if (r.get("summary", {}).get("success", 0) or 0) == 0]

    best_model = None
    best_accuracy = -1.0
    for report_row in completed:
        accuracy = report_row.get("summary", {}).get("accuracy")
        if isinstance(accuracy, (int, float)) and accuracy > best_accuracy:
            best_accuracy = float(accuracy)
            best_model = report_row.get("model")

    document.add_heading("Wnioski", level=2)
    document.add_paragraph(f"Modele z pełnym wykonaniem (success > 0): {len(completed)}")
    document.add_paragraph(f"Modele z błędami (success = 0): {len(failed)}")
    if best_model is not None:
        document.add_paragraph(
            f"Najlepszy model wg accuracy: {best_model} ({best_accuracy * 100:.1f}%)."
        )
    else:
        document.add_paragraph("Nie wyznaczono najlepszego modelu (brak poprawnie wykonanych predykcji).")

    if failed:
        document.add_heading("Modele niewykonane poprawnie", level=3)
        for failed_row in failed:
            model = failed_row.get("model", "-")
            errors = failed_row.get("summary", {}).get("errors", "-")
            document.add_paragraph(f"- {model}: errors={errors}")

    document.add_paragraph("")
    document.add_paragraph(f"Plik źródłowy JSON: {input_path}")
    document.add_paragraph(f"Wygenerowano: {datetime.now().isoformat()}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"Input report not found: {input_path}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if args.output:
        output_path = Path(args.output).resolve()
    else:
        output_path = input_path.parent / f"final_batch_report_{timestamp}.docx"

    report = load_report(input_path)
    build_doc(report, input_path, output_path)
    print(f"DOCX report saved to: {output_path}")


if __name__ == "__main__":
    main()
